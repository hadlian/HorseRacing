"""
BRIS Summary Handicap Report generator — Dennis format.

Takes the score_card() output dict and writes a .docx matching Dennis's
CDX0529_BRIS_Summary_Report.docx layout exactly.

Entry point:
    generate_bris_summary(results, card_name, out_path)

    results   — dict returned by comparemodels_engine.score_card()
                {race_num: {"ranked_horses": [...], "category_picks": {...}, ...}}
    card_name — e.g. "CDX0529"
    out_path  — absolute path for the output .docx file
"""

from docx import Document
from docx.shared import Pt, RGBColor
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# Display order + labels matching Dennis's format exactly
CATEGORY_ORDER = [
    ("Avg Speed",      "Top 3 Average Speeds"),
    ("Distance Speed", "Top 3 Distance Speeds"),
    ("Best Speed",     "Top 3 Best Speed"),
    ("Prime Power",    "Top 3 Prime Power"),
    ("Avg Class",      "Top 3 Average Class"),
    ("Jockey Rating",  "Top 3 Jockeys"),
    ("Trainer Rating", "Top 3 Trainers"),
    ("Earnings",       "Top 3 Earnings"),
]

# Max horses to show in Consensus Leaders line
CONSENSUS_LEADERS_MAX = 5
# Threshold for rollup tables
POSITIVE_FACTORS_MIN = 3


def _set_para_spacing(para, before: int = 0, after: int = 0):
    """Set paragraph spacing (in points * 914, i.e. twips) before/after."""
    pPr = para._p.get_or_add_pPr()
    spacing = OxmlElement("w:spacing")
    spacing.set(qn("w:before"), str(before * 20))
    spacing.set(qn("w:after"),  str(after  * 20))
    pPr.append(spacing)


def _add_line(doc, label: str, value: str, bold_label: bool = False):
    """Add a single 'Label: value' paragraph."""
    para = doc.add_paragraph()
    _set_para_spacing(para, before=0, after=0)
    lbl_run = para.add_run(label + ": ")
    lbl_run.bold = bold_label
    para.add_run(value)
    return para


def _build_category_line(doc, label: str, picks: list):
    """
    Add a category top-3 line.
    The top pick gets bold formatting when the underline rule fires
    (matches Dennis's bold marker for __pgm__ in his txt schema).
    Shows '—' when no picks exist for the category (e.g. all distance speeds are 0).
    """
    para = doc.add_paragraph()
    _set_para_spacing(para, before=0, after=0)
    para.add_run(label + ": ")

    if not picks:
        para.add_run("—")
    else:
        for i, pick in enumerate(picks[:3]):
            if i > 0:
                para.add_run(" - ")
            run = para.add_run(str(pick["pgm"]))
            if pick.get("underlined"):
                run.bold = True
    return para


def _build_composite_line(doc, ranked: list):
    """
    Composite Scores: pgm=score | pgm=score ...  (all horses, desc)
    """
    parts = " | ".join(f'{h["pgm"]}={h["composite"]}' for h in ranked)
    _add_line(doc, "Composite Scores", parts)


def _add_table(doc, rows_data: list):
    """
    Add a 3-column table: Horse Number | Horse Name | Positive Factors
    rows_data: list of (pgm, name, factor_count)
    """
    tbl = doc.add_table(rows=1, cols=3)
    tbl.style = "Table Grid"

    hdr_cells = tbl.rows[0].cells
    hdr_cells[0].text = "Horse Number"
    hdr_cells[1].text = "Horse Name"
    hdr_cells[2].text = "Positive Factors"
    for cell in hdr_cells:
        for run in cell.paragraphs[0].runs:
            run.bold = True

    for pgm, name, count in rows_data:
        row = tbl.add_row().cells
        row[0].text = str(pgm)
        row[1].text = str(name)
        row[2].text = str(count)


def generate_bris_summary(results: dict, card_name: str, out_path: str) -> None:
    """
    Generate BRIS Summary Handicap Report matching Dennis's format.

    results   — {race_num: score_race dict} from comparemodels_engine.score_card()
    card_name — e.g. "CDX0529" (used in title and filename)
    out_path  — full path to write the .docx
    """
    doc = Document()

    # Remove default paragraph spacing from Normal style
    style = doc.styles["Normal"]
    style.paragraph_format.space_before = Pt(0)
    style.paragraph_format.space_after  = Pt(0)

    # ── Title block ──────────────────────────────────────────────────────────
    title_para = doc.add_paragraph()
    title_run  = title_para.add_run("BRIS Summary Handicap Report")
    title_run.bold = True
    title_run.font.size = Pt(14)

    sub_para = doc.add_paragraph()
    sub_para.add_run(
        f"Generated from {card_name} using the BRIS Summary Handicap System."
        # Note: BRIS Top Pick (CM-4) and Jockey/Trainer field-position fixes are pending
        # engine work; consensus counts and composites will differ from Dennis's until resolved.
    )

    # ── Per-race blocks ───────────────────────────────────────────────────────
    for race_num in sorted(results.keys()):
        r       = results[race_num]
        ranked  = r.get("ranked_horses", [])
        cat_px  = r.get("category_picks", {})

        # Blank line + Race header
        doc.add_paragraph()
        race_para = doc.add_paragraph()
        race_run  = race_para.add_run(f"Race {race_num}")
        race_run.bold = True
        _set_para_spacing(race_para, before=6, after=2)

        # Category top-3 lines
        for cat_key, cat_label in CATEGORY_ORDER:
            picks = cat_px.get(cat_key, [])
            _build_category_line(doc, cat_label, picks)

        # Missing data notes (e.g. no distance figures, first-time starters)
        notes = r.get("missing_notes", [])
        if notes:
            note_para = doc.add_paragraph()
            _set_para_spacing(note_para, before=2, after=0)
            note_run = note_para.add_run("Note: " + " | ".join(notes))
            note_run.italic = True
            note_run.font.size = Pt(8)

        # BRIS Top Pick (CM-4 deferred — show pgm if is_bris_pick, else dash)
        bris_pick = next(
            (h["pgm"] for h in ranked if h.get("is_bris_pick")), "—"
        )
        _add_line(doc, "BRIS Top Pick", str(bris_pick))

        # Consensus Leaders (top 5 by consensus_count, tiebreak by composite)
        leaders = sorted(
            [h for h in ranked if h["consensus_count"] > 0],
            key=lambda h: (-h["consensus_count"], -h["composite"])
        )[:CONSENSUS_LEADERS_MAX]
        cons_str = " - ".join(
            f'{h["pgm"]}({h["consensus_count"]})' for h in leaders
        ) if leaders else "—"
        _add_line(doc, "Consensus Leaders", cons_str)

        # Dominant
        dominant = [h["pgm"] for h in ranked if h.get("is_dominant")]
        _add_line(doc, "Dominant", " - ".join(map(str, dominant)) if dominant else "—")

        # Pace leaders
        _add_line(doc, "Early Pace", str(r.get("early_pace_leader") or "—"))
        _add_line(doc, "Late Pace",  str(r.get("late_pace_leader")  or "—"))

        # Overlay Watch
        overlays = [h["pgm"] for h in ranked if h.get("is_overlay")]
        _add_line(doc, "Overlay Watch", " - ".join(map(str, overlays)) if overlays else "—")

        # A / B / C tiers — Dennis's format: A=rank1, B=ranks2-4 (max 3), C=ranks5-7 (max 3)
        tiers: dict = {"A": [], "B": [], "C": []}
        for h in ranked:
            t = h.get("tier", "C")
            tiers.setdefault(t, []).append(h["pgm"])
        for tier_label in ("A", "B", "C"):
            horses = tiers.get(tier_label, [])[:3]   # cap each tier at 3
            _add_line(doc, tier_label, " - ".join(map(str, horses)) if horses else "—")

        # Composite Scores (all horses, desc)
        _build_composite_line(doc, ranked)

    # ── End-of-card rollup ────────────────────────────────────────────────────
    doc.add_paragraph()
    rollup_hdr = doc.add_paragraph()
    run = rollup_hdr.add_run("HORSES WITH 3+ POSITIVE FACTORS")
    run.bold = True
    _set_para_spacing(rollup_hdr, before=8, after=4)

    for race_num in sorted(results.keys()):
        r      = results[race_num]
        ranked = r.get("ranked_horses", [])

        race_hdr = doc.add_paragraph(f"RACE {race_num} – HORSES WITH 3+ POSITIVE FACTORS")
        _set_para_spacing(race_hdr, before=6, after=2)

        qualifiers = sorted(
            [h for h in ranked if h["consensus_count"] >= POSITIVE_FACTORS_MIN],
            key=lambda h: (-h["consensus_count"], -h["composite"])
        )

        rows_data = [(h["pgm"], h["name"], h["consensus_count"]) for h in qualifiers]
        _add_table(doc, rows_data)

    doc.save(out_path)
